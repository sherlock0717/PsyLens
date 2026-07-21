#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the public PsyLens project brief from data/public."""
from __future__ import annotations

import argparse
import csv
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SAMPLES = ROOT / "data" / "public" / "samples_public.csv"
DEFAULT_EVIDENCE = ROOT / "data" / "public" / "evidence_public.csv"
DEFAULT_OUTPUT = ROOT / "docs" / "files" / "PsyLens_project_brief.docx"

TOPIC_CN = {
    "balance": "平衡与数值",
    "matchmaking": "匹配与对局分配",
    "event_design": "活动与玩法设计",
    "progression": "成长与养成",
    "community_conflict": "社区冲突与氛围",
    "communication_transparency": "沟通与透明度",
    "rewards": "奖励与产出",
    "new_player_onboarding": "新手引导与体验",
    "other_uncertain": "其他 / 不确定",
}
MECH_CN = {
    "competence_frustration": "胜任受挫",
    "fairness_threat": "公平受损",
    "trust_communication_gap": "信任与沟通落差",
    "belonging_drop": "归属感下降",
    "norm_safety_risk": "规范与安全风险",
    "uncertain": "不确定",
}
PLATFORM_CN = {"Bili": "B 站", "NGA": "NGA", "Tieba": "贴吧"}
NAVY = "12304A"
BLUE = "2176A3"
PALE = "EDF5F9"
GRAY = "F2F4F6"
TEXT = RGBColor(31, 47, 61)
MUTED = RGBColor(91, 108, 120)
ACCENT = RGBColor(33, 118, 163)
WHITE = RGBColor(255, 255, 255)


def read_csv(path: Path):
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def pct(value, total):
    return f"{value / total * 100:.1f}%" if total else "0.0%"


def font(run, size=10.5, bold=False, color=TEXT, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = tc_pr.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        tc_pr.append(element)
    element.set(qn("w:fill"), fill)


def border(cell, color=BLUE, size=7):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def row_rule(row, repeat=False):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if repeat:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def setup(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 30, RGBColor(18, 48, 74)),
        ("Heading 1", 20, RGBColor(18, 48, 74)),
        ("Heading 2", 15, RGBColor(33, 91, 126)),
        ("Heading 3", 12, RGBColor(33, 91, 126)),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(18 if name == "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("PsyLens｜社区反馈分析与可靠性评测"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("PsyLens  ·  "), size=8.5, color=MUTED)
    page_field(footer)


def paragraph(doc, text, size=10.5, color=TEXT, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    font(p.add_run(text), size=size, color=color, bold=bold)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item))


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        font(p.add_run(item))


def table(doc, headers, rows, widths=None, size=8.7):
    result = doc.add_table(rows=1, cols=len(headers))
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.style = "Table Grid"
    result.autofit = widths is None
    header = result.rows[0]
    row_rule(header, repeat=True)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(str(value)), size=size, bold=True, color=WHITE)
        if widths:
            cell.width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        row = result.add_row()
        row_rule(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            shade(cell, "FFFFFF" if row_index % 2 == 0 else GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), size=size)
            if widths:
                cell.width = Cm(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return result


def callout(doc, title, text, warning=False):
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    shade(cell, "FFF4DF" if warning else PALE)
    border(cell)
    p = cell.paragraphs[0]
    font(p.add_run(title + "\n"), size=10.5, bold=True, color=ACCENT)
    font(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    font(p.add_run("PSYLENS"), size=13, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(12)
    font(p.add_run("社区反馈分析与\n可靠性评测"), size=31, bold=True, color=RGBColor(18, 48, 74))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(25)
    font(p.add_run("海克斯大乱斗多平台案例｜公开发布版项目说明"), size=15, color=ACCENT)
    box = doc.add_table(rows=1, cols=1)
    cell = box.cell(0, 0)
    shade(cell, PALE)
    border(cell, size=10)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("360 条公开样本  ·  927 条证据单元  ·  三个平台等额抽样  ·  文本回溯率 100%"), size=11.5, bold=True, color=RGBColor(18, 73, 105))
    doc.add_paragraph("")
    for label, value in [
        ("项目定位", "社区反馈分析案例、可复现数据处理流程、心理学双层编码、可靠性评测与自动校准原型"),
        ("公开产物", "脱敏数据、编码手册、分析页面、离线 Demo 与复现脚本"),
        ("版本日期", "2026 年 7 月"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        font(p.add_run(label + "  "), size=10, bold=True, color=MUTED)
        font(p.add_run(value), size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    font(p.add_run("Copyright © 2026 Sherlock0717. All rights reserved."), size=8.5, color=MUTED)
    doc.add_page_break()


def compute(samples, evidence):
    specific = Counter(row["surface_topic"] for row in evidence if row["surface_topic"] != "other_uncertain")
    assigned = Counter(row["mechanism_label"] for row in evidence if row["mechanism_label"] != "uncertain")
    mechanisms = Counter(row["mechanism_label"] for row in evidence)
    cross = Counter((row["surface_topic"], row["mechanism_label"]) for row in evidence)
    density = Counter(row["sample_id"] for row in evidence)
    platform = {}
    for key in sorted({row["platform_source"] for row in samples}):
        platform_samples = [row for row in samples if row["platform_source"] == key]
        platform_evidence = [row for row in evidence if row["platform_source"] == key]
        platform[key] = {
            "samples": len(platform_samples),
            "evidence": len(platform_evidence),
            "density": len(platform_evidence) / len(platform_samples),
            "topics": Counter(row["surface_topic"] for row in platform_evidence if row["surface_topic"] != "other_uncertain"),
            "mechanisms": Counter(row["mechanism_label"] for row in platform_evidence if row["mechanism_label"] != "uncertain"),
        }
    densities = [density.get(row["sample_id"], 0) for row in samples]
    return specific, assigned, mechanisms, cross, platform, densities


def build(samples_path: Path, evidence_path: Path, output: Path):
    samples = read_csv(samples_path)
    evidence = read_csv(evidence_path)
    specific, assigned, mechanisms, cross, platforms, densities = compute(samples, evidence)
    n_evidence = len(evidence)
    specific_n = sum(specific.values())
    assigned_n = sum(assigned.values())
    date_missing = sum(1 for row in samples if not row.get("date", "").strip())
    no_evidence = sum(1 for value in densities if value == 0)

    doc = Document()
    setup(doc)
    cover(doc)

    doc.add_heading("执行摘要", level=1)
    paragraph(doc, "PsyLens 是一个社区反馈分析案例，同时提供可复现的数据处理流程、心理学启发的双层编码方法、可靠性评测工具和可接入模型的自动校准原型。真实模型运行结果、模型之间的一致性比较和人工金标准不属于当前版本的交付范围。")
    paragraph(doc, "PsyLens 将社区公开讨论整理成可回溯的证据单元，并用表层话题与体验机制两套编码描述问题位置和玩家体验方向。当前案例覆盖 NGA、贴吧与 B 站，每个平台等额保留 120 条样本。")
    paragraph(doc, "公开数据包含 360 条样本和 927 条证据，平均每条样本切出 2.575 条证据。样本 ID、证据 ID、父样本关联、平台字段与文本定位检查均通过；公开文件未检出 URL，规范化后未发现文本重复组。")
    table(doc, ["指标", "当前值", "说明"], [
        ["样本", "360", "三个平台各 120 条"],
        ["证据", "927", "平均 2.575 条 / 样本"],
        ["具体话题", "441", "平衡与数值占 61.2%"],
        ["已分配机制", "441", "胜任受挫占 69.8%"],
        ["机制不确定", "486", "占全部证据 52.4%"],
        ["结构回溯", "100%", "孤立证据 0，平台错配 0"],
    ], widths=[4.0, 3.0, 8.5], size=9)
    callout(doc, "核心判断", "主要具体话题集中在平衡与数值，主要已分配体验机制集中在胜任受挫；匹配与对局分配构成第二大具体话题。机制不确定率较高，后续重点应放在上下文补充、标签边界和双人复编码。")

    doc.add_heading("1. 项目问题与交付目标", level=1)
    paragraph(doc, "社区反馈具有文本短、情绪强、上下文分散和平台表达习惯不同等特点。单纯统计关键词容易混合不同语境，也难以判断结论能否回到原始表达。项目围绕三个问题组织流程：")
    numbered(doc, [
        "玩家正在集中讨论哪些具体问题？",
        "这些问题呈现了怎样的体验机制与归因方向？",
        "从文本到结论的每一步是否可核查、可解释并能够复现？",
    ])
    paragraph(doc, "交付结果包括公开脱敏样本、证据表、编码手册、心理学框架、分析页面、评测定义、离线 Demo 与本项目说明。")

    doc.add_heading("2. 数据设计与样本结构", level=1)
    doc.add_heading("2.1 平台与抽样", level=2)
    paragraph(doc, "当前案例使用三个社区平台，每个平台各保留 120 条样本。等额抽样用于确保平台都能进入分析，适合比较表达结构与编码表现；平台样本数不用于推断真实讨论规模。")
    platform_rows = [[PLATFORM_CN[key], values["samples"], values["evidence"], f"{values['density']:.3f}"] for key, values in ((key, platforms[key]) for key in ["Bili", "NGA", "Tieba"])]
    table(doc, ["平台", "样本数", "证据数", "每样本证据"], platform_rows, widths=[4.0, 3.0, 3.0, 4.0], size=9)
    paragraph(doc, f"样本文本长度中位数为 61 字，平均 84.03 字；90% 的样本文本不超过 164 字。{date_missing} 条样本日期为空，因此当前项目不进行时间趋势推断。")
    doc.add_heading("2.2 样本与证据的区别", level=2)
    paragraph(doc, "样本层保存一条完整反馈，证据层保存其中可独立判断的最小片段。一条样本可以产生零条、一条或多条证据。当前证据密度中位数为 2，90 分位为 5，最大值为 21；有 2 条样本未切出证据。")
    callout(doc, "为什么需要证据层", "完整反馈适合理解语境，证据单元适合编码和核对。两层同时保留，可以减少把复杂文本强行归入单一标签，并让数量结论回到具体原文。")

    doc.add_heading("3. 数据清洗与证据构建", level=1)
    paragraph(doc, "数据处理以保留原意、统一结构、降低公开风险和维持回溯关系为目标。完整操作分为以下八步：")
    table(doc, ["步骤", "操作"], [
        ["候选讨论登记", "按平台、主题和时间窗形成候选清单，记录主帖或回复类型。"],
        ["页面与回复采集", "保存原始响应与必要上下文，登录态和缓存留在本地环境。"],
        ["规则预清洗", "删除回复头、引用残留、图片路径和多余空白，移除空文本。"],
        ["内容筛选", "排除纯标记、极短噪声、离题文本和缺少体验信息的纯攻击表达。"],
        ["平台等额抽样", "三个平台各保留 120 条，并保留时间窗、粗主题与回复类型字段。"],
        ["公开脱敏", "删除来源链接、账号和平台内部定位字段，屏蔽联系方式与文本内链接。"],
        ["证据切分与编码", "按标点切分独立片段，记录话题、机制、编码来源与纳入状态。"],
        ["结构与分布审计", "检查 ID、父样本、文本定位、重复、空值、证据密度和标签分布。"],
    ], widths=[4.2, 11.3], size=8.8)
    doc.add_heading("3.1 公开字段最小化", level=2)
    paragraph(doc, "公开样本只保留稳定编号、平台类别、平台内顺序、时间窗、抽样粗主题、回复类型、可选日期、公开文本与迁移状态。旧公开数据同时保存内容相同的 raw_text 与 public_text，当前版本删除冗余 raw_text；47 条空表层话题统一写为 other_uncertain。")
    doc.add_heading("3.2 证据切分规则", level=2)
    bullets(doc, [
        "按中英文句号、问号、感叹号、分号与换行切分；",
        "切分后清理首尾空白，默认少于 6 个字符的片段不单独进入候选；",
        "不跨样本拼接，不补写省略信息，不改变原句立场；",
        "每条证据文本必须能够在父样本公开文本中逐字定位；",
        "信息不足时保留暂时不确定类别。",
    ])

    doc.add_heading("4. 心理学分析框架", level=1)
    paragraph(doc, "项目使用双层编码：表层话题记录文本正在讨论的具体问题，体验机制记录文本呈现的心理体验方向。这样既保留产品问题的具体位置，也形成跨话题比较的解释维度。")
    table(doc, ["机制", "理论线索", "文本中的操作性含义"], [
        ["胜任受挫", "自我决定理论中的胜任需要", "投入难以转化为有效行动、任务完成或可理解的进步"],
        ["公平受损", "公平判断与组织公正", "将结果归因于规则、匹配、资源或数值分配不合理"],
        ["信任与沟通落差", "组织信任与程序透明", "说明、回应、承诺或反馈闭环不足"],
        ["归属感下降", "社群归属与共同体感", "与游戏、社区或身份群体出现疏离"],
        ["规范与安全风险", "群体规范与心理安全", "外挂、辱骂、骚扰、误封或举报问题"],
        ["不确定", "测量中的保留类别", "文本过短、依赖上下文或主机制无法判定"],
    ], widths=[3.5, 4.5, 7.5], size=8.5)
    paragraph(doc, "机制标签用于组织文本证据与提出待验证问题，不用于推断玩家稳定人格、临床状态或真实能力。项目采用固定纳入、排除和相邻标签裁决规则，并保留 uncertain 以控制过度解释。")

    doc.add_heading("5. 分析结果", level=1)
    doc.add_heading("5.1 具体话题分布", level=2)
    table(doc, ["具体话题", "证据数", "占具体话题证据"], [[TOPIC_CN[key], count, pct(count, specific_n)] for key, count in specific.most_common()], widths=[8.0, 3.0, 4.5], size=9)
    paragraph(doc, f"在 {specific_n} 条能够分配具体话题的证据中，平衡与数值为 {specific['balance']} 条，占 {pct(specific['balance'], specific_n)}；匹配与对局分配为 {specific['matchmaking']} 条，占 {pct(specific['matchmaking'], specific_n)}。两类合计覆盖 79.8% 的具体话题证据。")
    doc.add_heading("5.2 体验机制分布", level=2)
    table(doc, ["已分配机制", "证据数", "占已分配机制证据"], [[MECH_CN[key], count, pct(count, assigned_n)] for key, count in assigned.most_common()], widths=[8.0, 3.0, 4.5], size=9)
    paragraph(doc, f"在 {assigned_n} 条能够明确分配机制的证据中，胜任受挫为 {assigned['competence_frustration']} 条，占 {pct(assigned['competence_frustration'], assigned_n)}；公平受损为 {assigned['fairness_threat']} 条，占 {pct(assigned['fairness_threat'], assigned_n)}。")
    callout(doc, "重要质量信号", f"另有 {mechanisms['uncertain']} 条证据的机制为 uncertain，占全部证据 {pct(mechanisms['uncertain'], n_evidence)}。高不确定率提示短文本、上下文缺失和标签边界仍是主要风险。该指标应与 163 条证据纳入提醒分开理解。", warning=True)
    doc.add_heading("5.3 话题 × 机制交叉", level=2)
    combinations = [((topic, mechanism), count) for (topic, mechanism), count in cross.items() if topic != "other_uncertain" and mechanism != "uncertain"]
    combinations.sort(key=lambda item: (-item[1], item[0]))
    table(doc, ["组合", "证据数", "占全部证据"], [[f"{TOPIC_CN[topic]} × {MECH_CN[mechanism]}", count, pct(count, n_evidence)] for (topic, mechanism), count in combinations[:10]], widths=[9.2, 2.8, 3.5], size=8.8)
    paragraph(doc, "平衡 × 胜任受挫是数量最高的具体组合，说明角色、装备和机制强度经常与“难以发挥作用”同时出现。平衡 × 公平受损与匹配 × 公平受损提示部分玩家把负面结果归因于规则或分配方式。沟通透明 × 信任落差数量较少，但体验类型对应关系清晰。")
    doc.add_heading("5.4 平台描述性比较", level=2)
    platform_detail = []
    for key in ["Bili", "NGA", "Tieba"]:
        values = platforms[key]
        top_topics = "、".join(TOPIC_CN[item] for item, _ in values["topics"].most_common(3))
        top_mechanisms = "、".join(MECH_CN[item] for item, _ in values["mechanisms"].most_common(2))
        platform_detail.append([PLATFORM_CN[key], values["samples"], values["evidence"], f"{values['density']:.2f}", top_topics, top_mechanisms])
    table(doc, ["平台", "样本", "证据", "证据 / 样本", "主要具体话题", "主要机制"], platform_detail, widths=[2.0, 1.4, 1.5, 2.0, 4.5, 4.0], size=8.1)
    paragraph(doc, "NGA 的证据密度最高，B 站最低。该差异可能来自原文长度、切分粒度和表达风格。NGA 与贴吧主要保留历史 AI 编码，B 站使用离线规则提案，编码来源差异也会影响分布，因此这里只做描述性比较。")
    doc.add_heading("5.5 证据示例", level=2)
    table(doc, ["证据 ID", "话题", "机制", "公开文本"], [
        ["NGA_0021_U01", "匹配与对局分配", "公平受损", "匹配机制就导致了一旦你连胜或者连败，要把你近期胜率“拉”回五十接下来就会反向连胜连败"],
        ["TIEBA_0004_U02", "新手引导与体验", "胜任受挫", "通常是莫甘娜，我新手目前英雄正在熟悉，怕坑，所以会选一些弱一点，但熟悉的英雄，为了避免伤害没打多少，还抗不了伤害，就会多出功能装"],
        ["NGA_0008_U01", "平衡与数值", "公平受损", "现在怒气增长吃终极技能急速，残疫、钢壁甚至猎魔人都能增加怒气获取，但偏偏不吃技能急速，公理圆弧效果也不能直接涨一截怒气。"],
    ], widths=[2.8, 3.2, 2.6, 7.0], size=7.9)

    doc.add_heading("6. 可靠性评测", level=1)
    paragraph(doc, "评测覆盖结构完整性、清洗与隐私、编码可用性、分析支撑和运行复现五个维度。硬性错误阻止发布，诊断指标用于解释风险与规划返修。")
    table(doc, ["指标", "当前值", "判定", "说明"], [
        ["样本 ID 唯一率", "100%", "通过", "无重复样本编号"],
        ["证据 ID 唯一率", "100%", "通过", "无重复证据编号"],
        ["父样本关联率", "100%", "通过", "孤立证据为 0"],
        ["文本回溯率", "100%", "通过", "927 条证据均可定位"],
        ["平台错配", "0", "通过", "证据与父样本平台一致"],
        ["公开 URL 命中", "0", "通过", "公开数据未检出链接"],
        ["样本 / 证据重复组", "0 / 0", "通过", "规范化文本扫描"],
        ["无证据样本", str(no_evidence), "记录", "占样本 0.6%"],
        ["机制不确定率", pct(mechanisms['uncertain'], n_evidence), "重点诊断", "需要改进上下文与标签边界"],
        ["缺失日期", str(date_missing), "限制时间分析", "日期属于可选字段"],
    ], widths=[4.2, 2.4, 3.2, 5.7], size=8.4)
    doc.add_heading("6.1 编码一致性评测方法（扩展）", level=2)
    paragraph(doc, "编码一致性评测从平台、主要话题、主要机制和不确定类别中分层抽样，由两名编码者独立判断证据纳入、表层话题与机制标签。记录独立结果、争议类型、裁决理由和手册修订，并计算适用于名义变量的 Krippendorff’s alpha。Krippendorff’s alpha 是一种衡量多名编码者判断一致程度的统计量。")
    paragraph(doc, "一致性指标需要与混淆矩阵、争议样本和返修成本一起解释。较高一致性只能说明编码者使用同一规则时结果稳定，不能单独证明体验类型划分完整或产品结论有效。")

    doc.add_heading("7. 文案质量检查", level=1)
    paragraph(doc, "公开文案由 tools/lint_public_copy.py 自动检查，覆盖 README、页面、方法文档以及 Python 中面向读者的文字。检查器识别模板化宣传词、负面框架句式、过长句子、页面中的内部工程状态和未解释的英文缩写。检查器设有发布门槛（一道自动关卡）：只有文案检查达到设定级别，页面和文档才允许发布。仓库在 Ubuntu 与 Windows 两套环境运行持续集成，包含编译检查、Ruff 静态检查、测试套件、文案质量门槛，以及一次本地固定示例校准流程的冒烟运行与公开字段安全检查。")

    doc.add_heading("8. 自动校准工具", level=1)
    paragraph(doc, "仓库提供一套编码复检与争议分析工具，用于观察标签在多次独立判断下是否稳定。它是可接入模型的自动校准原型，当前公开版本不发布模型成绩。")
    doc.add_heading("8.1 分层抽样与三路复检", level=2)
    paragraph(doc, "抽样从公开证据中取 300 条主样本，另加 30 条重测样本，按平台、话题、机制、编码来源、纳入状态和文本长度分层。重测样本用于观察同一条证据在重复判断下是否得到一致结果。两部分合并后统一编号并重新排序，公开样本只保留复检所需的脱敏字段，来源、平台、当前标签和重测关系写入私有映射文件。")
    paragraph(doc, "三名独立 Reviewer 使用不同的判断结构，但共用同一套标签集合和定义：Reviewer A 严格依据编码手册分类；Reviewer B 先概括主要诉求再匹配标签；Reviewer C 先排除相邻标签再做判断。三名 Reviewer 互不查看彼此结果。共识分析汇总一致标签与争议标签，争议项转化为编码手册的改进提案。")
    doc.add_heading("8.2 运行方式与扩展入口", level=2)
    paragraph(doc, "工具提供两种运行方式。本地固定示例模式按关键词生成确定性示例输出，用来验证数据流、输出格式和统计流程；它只用于验证流程，不是真实模型校准结果。OpenAI-compatible 接口作为扩展入口保留，可在本地配置环境变量后接入外部模型；当前公开版本不包含真实模型运行结果，也不发布模型成绩。")
    callout(doc, "结果性质说明", "本地固定示例只用于验证流程；当前版本不发布模型成绩；自动结果定位为参考，不等同于人工金标准。")

    doc.add_heading("9. 产品方向与验证设计", level=1)
    paragraph(doc, "当前结果适合形成产品假设，后续通过独立研究验证。以下方向根据高频话题、主要机制与证据示例整理：")
    table(doc, ["方向", "证据线索", "产品假设", "验证指标"], [
        ["角色与强化适配信息", "平衡 × 胜任受挫", "在选择阶段提供强化适配、阵容缺口与风险提示", "A/B 对比选择后退出率、低贡献局比例、相关负面反馈"],
        ["匹配预期与解释", "匹配 × 公平受损", "提供可理解的匹配范围、队列状态和连胜连败解释", "访谈理解度；上线前后公平归因反馈与中途退出"],
        ["新手安全选择路径", "新手引导 × 胜任受挫", "按熟悉度、操作难度和团队作用提供推荐与功能装提示", "新手任务成功率、首局贡献、二次参与率"],
        ["改动沟通闭环", "沟通透明 × 信任落差", "公告补充改动目的、影响范围、反馈入口与后续复盘", "公告理解测试、反馈重复率、信任量表短项"],
        ["社区规范与处置反馈", "社区冲突 × 规范安全", "优化举报分类、处置回执和高风险互动提示", "举报完成率、重复举报、辱骂暴露与社区安全感"],
    ], widths=[3.1, 3.2, 5.2, 4.0], size=7.9)
    paragraph(doc, "验证时应预先定义目标用户、干预内容、主要指标、观察窗口和停止条件。反馈数量下降需要结合活跃度、任务完成、对局质量与留存指标，避免把沉默误判为体验改善。")

    doc.add_heading("10. 解释边界", level=1)
    bullets(doc, [
        "三个平台各 120 条属于等额抽样设计，不能代表平台真实讨论规模或总体玩家意见占比；",
        "证据数量受到文本长度和切分粒度影响，不能直接等同于问题强度；",
        "当前编码包含历史 AI 结果与离线规则提案，不同编码来源的分布不可视为完全等价；",
        "240 条样本缺少日期，当前分析不支持完整时间趋势判断；",
        "体验机制是文本编码维度，不构成个体心理诊断；",
        "产品方向需要通过访谈、问卷、行为日志或实验验证。",
    ])

    doc.add_heading("11. 复现与仓库入口", level=1)
    paragraph(doc, "公开数据规范化、统计汇总与 Demo 均可离线运行。")
    for command in [
        "python tools/normalize_public_dataset.py --source-dir data/public --output-dir artifacts/normalized_public",
        "python tools/summarize_public_analysis.py --public-dir data/public --output artifacts/public_analysis_summary.json",
        "python tools/run_demo.py --provider mock --output artifacts/demo/run",
        "python -m pytest demo/tests tests -q",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        font(p.add_run(command), name="Consolas", size=8.5, color=RGBColor(28, 78, 105))
    table(doc, ["路径", "内容"], [
        ["data/public/", "公开样本、证据、文件校验记录与字段说明"],
        ["docs/methodology/", "心理学框架、证据切分、话题与机制手册"],
        ["docs/evaluation/", "评测方法与当前结果"],
        ["pipeline/", "采集契约、配置与 Prompt 模板"],
        ["demo/", "离线端到端示例"],
        ["tools/", "数据规范化、统计汇总与 Demo 入口"],
    ], widths=[5.0, 10.5], size=9)

    doc.add_heading("参考文献", level=1)
    references = [
        "Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77–101.",
        "Colquitt, J. A. (2001). On the dimensionality of organizational justice. Journal of Applied Psychology, 86(3), 386–400.",
        "Deci, E. L., & Ryan, R. M. (2000). The “what” and “why” of goal pursuits. Psychological Inquiry, 11(4), 227–268.",
        "Krippendorff, K. (2018). Content Analysis: An Introduction to Its Methodology (4th ed.). SAGE.",
        "Mayer, R. C., Davis, J. H., & Schoorman, F. D. (1995). An integrative model of organizational trust. Academy of Management Review, 20(3), 709–734.",
        "McMillan, D. W., & Chavis, D. M. (1986). Sense of community: A definition and theory. Journal of Community Psychology, 14(1), 6–23.",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        font(p.add_run(reference), size=9.2)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main(argv=None):
    parser = argparse.ArgumentParser(description="Build the PsyLens public project brief")
    parser.add_argument("--samples", default=str(DEFAULT_SAMPLES))
    parser.add_argument("--evidence", default=str(DEFAULT_EVIDENCE))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    args = parser.parse_args(argv)
    output = build(Path(args.samples), Path(args.evidence), Path(args.output))
    print(output)


if __name__ == "__main__":
    main()
